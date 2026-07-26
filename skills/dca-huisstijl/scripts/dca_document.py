#!/usr/bin/env python3
"""
dca_document.py — genereert DOCX, HTML of PDF in DCA-huisstijl.

Gebruik:
    python dca_document.py --type pdf  --titel "Titel" --sub "ondertitel" \
        --in body.md --out /mnt/user-data/outputs/stuk.pdf
    python dca_document.py --type docx --titel "Titel" --in body.md --out stuk.docx
    python dca_document.py --type html --titel "Titel" --in body.md --out stuk.html

body.md = markdown-subset (zie MARKDOWN-SUBSET hieronder).

De PDF-tak schrijft standaard de bron-markdown naast de PDF weg
(<out>.src.md), zodat een volgende versie een edit is en geen herbouw.
Uitzetten met --no-source.

MARKDOWN-SUBSET
---------------
  # H1                     kop niveau 1 (Merriweather/serif, diepblauw)
  ## H2                    kop niveau 2 (+ gouden onderlijn)
  ### H3                   kop niveau 3
  - bullet                 opsomming
  1. genummerd             genummerde opsomming
  | a | b |                tabel; scheidingsrij |---|---| is optioneel
  > tekst                  CTA-blok onderaan (gouden rand, cream)
  !!! kern                 donkerblauw kernblok tot een lege regel
  !!! box                  wit kader met gouden rand tot een lege regel
  ::: cols                 twee (of meer) kolommen naast elkaar
  +++                      kolomscheiding binnen een ::: cols-blok
  :::                      sluit het kolommenblok
  <!--pb-->                pagina-einde
  **vet**  *cursief*  `code`  [tekst](url)  ==goud gemarkeerd==

Fonts: Merriweather/Inter zijn webfonts en zitten zelden lokaal.
Het script valt terug op Lora/DejaVu Serif (koppen) en DejaVu Sans (body).
Beide dekken Cyrillisch en Grieks; dat is bewust, want Merriweather-fallback
op een systeemserif brak eerder op niet-Latijnse tekens.
"""
import argparse
import html as _html
import os
import re
import subprocess
import sys

# ---------------------------------------------------------------- tokens
DIEPBLAUW = "1B2C5A"
DONKERBLAUW = "0F1C3F"
GOUD = "C9A227"
CREAM = "FAF7F0"
CREAM_ALT = "F2EDE1"
LIJN = "d8d2c4"
KOP_FONT = "Merriweather"
BODY_FONT = "Inter"
KOP_STACK = '"Merriweather", "Lora", "DejaVu Serif", Georgia, serif'
BODY_STACK = '"Inter", "DejaVu Sans", Helvetica, Arial, sans-serif'


def _hex_to_rgb(h):
    return tuple(int(h[i:i + 2], 16) for i in (0, 2, 4))


# ---------------------------------------------------------------- inline
# RAG-bolletjes. wkhtmltopdf heeft geen kleuren-emoji-font, dus 🔴🟠🟢 komen
# er monochroom uit. We vervangen ze door gekleurde spans; dan klopt de kleur
# gegarandeerd. Amber = DCA-goud, niet oranje (huisstijlregel).
RAG = {
    "\U0001F534": "#C0392B",   # rood
    "\U0001F7E0": "#" + GOUD,  # amber → DCA-goud
    "\U0001F7E1": "#" + GOUD,  # geel → DCA-goud
    "\U0001F7E2": "#2E7D32",   # groen
}


def _rag(t):
    for ch, kleur in RAG.items():
        t = t.replace(
            ch, f'<span style="color:{kleur};font-size:11pt;'
                f'line-height:0">&#9679;</span>')
    return t


def inline(t):
    """Inline-markdown → HTML. Escapet eerst, dus bron mag < en & bevatten."""
    t = _html.escape(t, quote=False)
    t = _rag(t)
    t = re.sub(r"`([^`]+)`", r'<code>\1</code>', t)
    t = re.sub(r"==([^=]+)==",
               f'<span style="color:#{GOUD};font-weight:bold">' + r"\1" + "</span>", t)
    t = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", t)
    t = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<i>\1</i>", t)
    t = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', t)
    return t


def _row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def _is_sep(line):
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


# ---------------------------------------------------------------- markdown
def md_to_html(md):
    out = []
    lines = md.splitlines()
    i = 0
    n = len(lines)

    while i < n:
        raw = lines[i]
        t = raw.rstrip()
        s = t.strip()

        if not s:
            i += 1
            continue

        # kolommenblok:  ::: cols  …  +++  …  :::
        if re.match(r"^:::\s*cols\s*$", s, re.I):
            i += 1
            cols, buf = [], []
            while i < n and not re.match(r"^:::\s*$", lines[i].strip()):
                if lines[i].strip() == "+++":
                    cols.append("\n".join(buf))
                    buf = []
                else:
                    buf.append(lines[i])
                i += 1
            i += 1  # sluitende :::
            cols.append("\n".join(buf))
            cols = [c for c in cols if c.strip()]
            if cols:
                w = round(100 / len(cols), 2)
                tds = "".join(
                    f'<td class="colcell" style="width:{w}%">{md_to_html(c)}</td>'
                    for c in cols)
                out.append(f'<table class="cols"><tr>{tds}</tr></table>')
            continue

        # pagina-einde
        if s in ("<!--pb-->", "<!-- pb -->"):
            out.append('<div class="pb"></div>')
            i += 1
            continue

        # kern- / box-blok
        m = re.match(r"^!!!\s*(kern|box)\s*$", s, re.I)
        if m:
            kind = m.group(1).lower()
            i += 1
            buf = []
            while i < n and lines[i].strip():
                buf.append(inline(lines[i].strip()))
                i += 1
            out.append(f'<div class="{kind}">' + "<br>".join(buf) + "</div>")
            continue

        # tabel
        if s.startswith("|") and s.endswith("|"):
            head = _row(s)
            i += 1
            if i < n and _is_sep(lines[i]):
                i += 1
            body = []
            while i < n and lines[i].strip().startswith("|"):
                body.append(_row(lines[i]))
                i += 1
            th = "".join(f"<th>{inline(c)}</th>" for c in head)
            trs = []
            for r in body:
                tds = "".join(f"<td>{inline(c)}</td>" for c in r)
                trs.append(f"<tr>{tds}</tr>")
            out.append(f"<table><tr>{th}</tr>{''.join(trs)}</table>")
            continue

        # blockquote → CTA
        if s.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(inline(lines[i].strip().lstrip(">").strip()))
                i += 1
            out.append('<div class="cta">' + "<br>".join(buf) + "</div>")
            continue

        # bullets
        if re.match(r"^[-*]\s+", s):
            items = []
            while i < n and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(inline(re.sub(r"^[-*]\s+", "", lines[i].strip())))
                i += 1
            out.append("<ul>" + "".join(f"<li>{x}</li>" for x in items) + "</ul>")
            continue

        # genummerd
        if re.match(r"^\d+[.)]\s+", s):
            items = []
            while i < n and re.match(r"^\d+[.)]\s+", lines[i].strip()):
                items.append(inline(re.sub(r"^\d+[.)]\s+", "", lines[i].strip())))
                i += 1
            out.append("<ol>" + "".join(f"<li>{x}</li>" for x in items) + "</ol>")
            continue

        # koppen
        if s.startswith("### "):
            out.append(f"<h3>{inline(s[4:])}</h3>")
            i += 1
            continue
        if s.startswith("## "):
            out.append(f"<h2>{inline(s[3:])}</h2>")
            i += 1
            continue
        if s.startswith("# "):
            out.append(f"<h1class>{inline(s[2:])}</h1class>".replace(
                "h1class", "h1"))
            i += 1
            continue

        # alinea
        buf = [s]
        i += 1
        while i < n and lines[i].strip() and not re.match(
                r"^(#|\||>|!!!|[-*]\s|\d+[.)]\s|<!--pb)", lines[i].strip()):
            buf.append(lines[i].strip())
            i += 1
        out.append("<p>" + inline(" ".join(buf)) + "</p>")

    return "\n".join(out)


# ---------------------------------------------------------------- css
CSS = f"""
@page {{ size: A4; margin: 16mm 14mm; }}
body {{ background:#{CREAM}; color:#1a1a1a; font-family:{BODY_STACK};
        font-size:9.6pt; line-height:1.45; margin:0; }}
h1 {{ font-family:{KOP_STACK}; color:#{DIEPBLAUW}; font-size:22pt; margin:0 0 2mm; }}
h2 {{ font-family:{KOP_STACK}; color:#{DIEPBLAUW}; font-size:13pt;
      margin:7mm 0 2mm; border-bottom:1.5px solid #{GOUD}; padding-bottom:1mm; }}
h3 {{ font-family:{KOP_STACK}; color:#{DONKERBLAUW}; font-size:10.5pt;
      margin:4mm 0 1.5mm; }}
p {{ margin:1.5mm 0 2.5mm; }}
.goldline {{ border:0; border-top:3px solid #{GOUD}; margin:0 0 3mm; }}
.sub {{ color:#{DONKERBLAUW}; font-size:9pt; margin:0 0 6mm; }}
table {{ border-collapse:collapse; width:100%; margin:2mm 0 3mm; }}
th {{ background:#{DIEPBLAUW}; color:#{CREAM}; text-align:left;
      padding:2mm 2.2mm; font-size:9pt; font-weight:bold; }}
td {{ padding:1.8mm 2.2mm; border-bottom:1px solid #{LIJN};
      vertical-align:top; font-size:9.3pt; }}
tr:nth-child(even) td {{ background:#{CREAM_ALT}; }}
ul, ol {{ margin:1.5mm 0 3mm; padding-left:5mm; }}
li {{ margin-bottom:1.2mm; }}
code {{ font-family:"DejaVu Sans Mono",monospace; font-size:8.8pt;
        background:#{CREAM_ALT}; padding:0 1mm; }}
a {{ color:#{DIEPBLAUW}; }}
.kern {{ background:#{DIEPBLAUW}; color:#{CREAM}; padding:3.5mm 4mm;
         border-left:5px solid #{GOUD}; margin:3mm 0 5mm; font-size:10pt; }}
.kern b {{ color:#E8C75A; }}
.box {{ background:#fff; border:1px solid #{GOUD}; padding:3mm 4mm; margin:3mm 0; }}
.cta {{ border-left:5px solid #{GOUD}; background:#{CREAM_ALT};
        padding:3.5mm 4mm; margin:6mm 0 0; font-size:9.5pt; }}
.cta b {{ color:#{DIEPBLAUW}; }}
.pb {{ page-break-before:always; }}
table.cols {{ margin:0; }}
table.cols td.colcell {{ border:0; background:none !important;
        padding:0 3mm 0 0; vertical-align:top; }}
table.cols h3 {{ margin-top:0; }}
.foot {{ margin-top:6mm; font-size:8pt; color:#6b6455;
         border-top:1px solid #{LIJN}; padding-top:2mm; }}
"""

FOOT = ("Delemarre Chess Academy · W. van Otterloolaan 7, Oegstgeest "
        "· schaaktrainer.nl")


def build_html(titel, sub, body_md, lang="nl"):
    return f"""<!DOCTYPE html>
<html lang="{lang}"><head><meta charset="utf-8">
<title>{_html.escape(titel)}</title><style>{CSS}</style></head><body>
<h1>{_html.escape(titel)}</h1>
<hr class="goldline">
{f'<p class="sub">{inline(sub)}</p>' if sub else ''}
{md_to_html(body_md)}
<p class="foot">{FOOT}</p>
</body></html>
"""


# ---------------------------------------------------------------- pdf
def build_pdf(titel, sub, body_md, out_path, lang="nl", keep_source=True):
    exe = None
    for cand in ("wkhtmltopdf",):
        try:
            subprocess.run([cand, "--version"], capture_output=True, check=True)
            exe = cand
            break
        except Exception:
            pass
    if exe is None:
        sys.exit("❌ wkhtmltopdf niet gevonden. apt-get install wkhtmltopdf")

    tmp = out_path + ".tmp.html"
    with open(tmp, "w", encoding="utf-8") as f:
        f.write(build_html(titel, sub, body_md, lang))

    cmd = [exe, "--enable-local-file-access", "--encoding", "utf-8",
           "--margin-top", "14mm", "--margin-bottom", "14mm",
           "--margin-left", "12mm", "--margin-right", "12mm",
           "--quiet", tmp, out_path]
    subprocess.run(cmd, check=True)
    os.remove(tmp)

    if keep_source:
        src = os.path.splitext(out_path)[0] + ".src.md"
        with open(src, "w", encoding="utf-8") as f:
            f.write(f"<!-- titel: {titel} -->\n")
            if sub:
                f.write(f"<!-- sub: {sub} -->\n")
            f.write(f"<!-- lang: {lang} -->\n\n")
            f.write(body_md)
        print(f"📝 Bron bewaard: {src}")

    print(f"✅ PDF opgeslagen: {out_path}")


# ---------------------------------------------------------------- docx
def build_docx(titel, body_md, out_path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.5)
        s.left_margin = s.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(titel)
    run.font.name = KOP_FONT
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*_hex_to_rgb(DIEPBLAUW))

    line = doc.add_paragraph()
    lrun = line.add_run("―" * 24)
    lrun.font.color.rgb = RGBColor(*_hex_to_rgb(GOUD))
    lrun.font.size = Pt(11)

    pending = []

    def flush_table():
        if not pending:
            return
        rows = [_row(x) for x in pending if not _is_sep(x)]
        pending.clear()
        if not rows:
            return
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        for ri, r in enumerate(rows):
            for ci, c in enumerate(r[:len(rows[0])]):
                cell = t.cell(ri, ci)
                cell.text = re.sub(r"\*\*|\*|`", "", c)
                if ri == 0:
                    for par in cell.paragraphs:
                        for rr in par.runs:
                            rr.font.bold = True
                            rr.font.color.rgb = RGBColor(*_hex_to_rgb(DIEPBLAUW))

    for raw in body_md.splitlines():
        t = raw.rstrip()
        s = t.strip()
        if s.startswith("|"):
            pending.append(s)
            continue
        flush_table()
        if not s or s.startswith("<!--") or s.startswith("!!!"):
            continue
        clean = re.sub(r"\*\*|\*|`", "", s)
        if s.startswith("### "):
            h = doc.add_paragraph(); r = h.add_run(clean[4:])
            r.font.name = KOP_FONT; r.font.bold = True; r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(*_hex_to_rgb(DONKERBLAUW))
        elif s.startswith("## "):
            h = doc.add_paragraph(); r = h.add_run(clean[3:])
            r.font.name = KOP_FONT; r.font.bold = True; r.font.size = Pt(15)
            r.font.color.rgb = RGBColor(*_hex_to_rgb(DIEPBLAUW))
        elif s.startswith("# "):
            h = doc.add_paragraph(); r = h.add_run(clean[2:])
            r.font.name = KOP_FONT; r.font.bold = True; r.font.size = Pt(19)
            r.font.color.rgb = RGBColor(*_hex_to_rgb(DIEPBLAUW))
        elif re.match(r"^[-*]\s+", s):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", clean), style="List Bullet")
        elif s.startswith(">"):
            doc.add_paragraph(clean.lstrip(">").strip(), style="Intense Quote")
        else:
            doc.add_paragraph(clean)
    flush_table()

    doc.save(out_path)
    print(f"✅ DOCX opgeslagen: {out_path}")


# ---------------------------------------------------------------- cli
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--type", choices=["docx", "pdf", "html"], default="docx")
    ap.add_argument("--titel", required=True)
    ap.add_argument("--sub", default="")
    ap.add_argument("--lang", default="nl")
    ap.add_argument("--in", dest="infile", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--no-source", action="store_true",
                    help="bron-markdown NIET naast de PDF wegschrijven")
    args = ap.parse_args()

    with open(args.infile, encoding="utf-8") as f:
        body = f.read()

    # front-matter uit een eerder bewaarde .src.md overnemen
    tit = re.search(r"<!--\s*titel:\s*(.+?)\s*-->", body)
    sub = re.search(r"<!--\s*sub:\s*(.+?)\s*-->", body)
    lan = re.search(r"<!--\s*lang:\s*(.+?)\s*-->", body)
    titel = args.titel or (tit.group(1) if tit else "")
    subtitel = args.sub or (sub.group(1) if sub else "")
    lang = args.lang if args.lang != "nl" else (lan.group(1) if lan else "nl")
    body = re.sub(r"<!--\s*(titel|sub|lang):.*?-->\n?", "", body)

    if args.type == "docx":
        build_docx(titel, body, args.out)
    elif args.type == "html":
        with open(args.out, "w", encoding="utf-8") as f:
            f.write(build_html(titel, subtitel, body, lang))
        print(f"✅ HTML opgeslagen: {args.out}")
    else:
        build_pdf(titel, subtitel, body, args.out, lang,
                  keep_source=not args.no_source)


if __name__ == "__main__":
    main()
